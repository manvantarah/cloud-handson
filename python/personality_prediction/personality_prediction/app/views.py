from django.shortcuts import render
from django.http import HttpResponseRedirect
from django.http import HttpResponseRedirect
from django.urls import reverse
from django.core.files.storage import FileSystemStorage
import os
import pickle
from docx import Document
from django.contrib import messages
import pandas as pd
import numpy as np
from sklearn import preprocessing
from sklearn import datasets, linear_model
from sklearn.metrics import mean_squared_error, r2_score
from sklearn import metrics
from sklearn.model_selection import train_test_split
from sklearn import neighbors

# Create your views here.
def index(request):
	return render(request,"index.html")

def candidate(request):
	if request.method == 'POST':
		global tfmax, tfres
		fullname = request.POST.get('name')
		age = request.POST.get('age')
		#print(age)
		file = request.FILES.get('document')
		document = Document(file)
		tables = document.tables

		for row in tables[0].rows:
			extract1 = row.cells[0].paragraphs[0].text+' : '+ row.cells[1].paragraphs[0].text +'\n'	
			#print(extract)

			dir_path = os.path.join('D:\\INTERNSHIP\\DJANGO\\personality_prediction\\media')
			with open(os.path.join(dir_path,fullname + ".txt"), 'a') as file:					
				file.write(extract1)

		for row in tables[1].rows:
			extract2 = row.cells[0].paragraphs[0].text 

		for row in tables[2].rows:
			extract3 = row.cells[0].paragraphs[0].text 
			
		docA = extract2.replace(',', '')
		docB = extract3.replace(',', '')
		doc= docA + " " + docB
		bow = doc.split(" ")
		wordDict = dict.fromkeys(bow, 0)
		for word in bow:
		    wordDict[word]+=1

		def computeTF(wordDict, bow):
		    tfDict = {}
		    bowCount = len(bow)
		    for word, count in wordDict.items():
		        tfDict[word] = count/float(bowCount)
		    return tfDict

		if "in" in wordDict:
		    del wordDict["in"]
		if "on" in wordDict:
		    del wordDict["on"]
		if "the" in wordDict:
		    del wordDict["the"]
		print(wordDict)
		tfBow = computeTF(wordDict, bow)
		print(tfBow)
		tfmax = max(tfBow, key=tfBow.get)
		print(tfmax)
		with open(os.path.join(dir_path,fullname + ".txt"), 'a') as file:									
			file.write("Candidate is good in: ") 
			file.write(tfmax)
		messages.success(request,"Your details are successfully submitted!")
	
	#html = "Please enter the credentials"
	return render(request, 'candidate.html')

def test(request):
	#candidate(request) 
	print(tfmax)
	global Openness, Conscientiousness, Extraversion, Aggreeableness, Neuroticism, age
	#print(age, fullname)
	if request.method=='POST':
		#Openness
		a= request.POST.get('a1')
		b= request.POST.get('a2')
		c= request.POST.get('a3')
		d= request.POST.get('a4')
		#conscientioness
		e= request.POST.get('b1')
		f= request.POST.get('b2')
		g= request.POST.get('b3')
		h= request.POST.get('b4')
		#extraversion
		i= request.POST.get('c1')
		j= request.POST.get('c2')
		k= request.POST.get('c3')
		l= request.POST.get('c4')
		#Aggreeableness
		m= request.POST.get('d1')
		n= request.POST.get('d2')
		o= request.POST.get('d3')
		p= request.POST.get('d4')
		#neuroticism
		q= request.POST.get('e1')
		r= request.POST.get('e2')
		s= request.POST.get('e3')
		t= request.POST.get('e4')

		Open = int(a)+int(b)+int(c)+int(d)
		print(Open)
		Cons = int(e)+int(f)+int(g)+int(h)
		print(Cons)
		Extra = int(i)+int(j)+int(k)+int(l)
		print(Extra)
		Aggre = int(m)+int(n)+int(o)+int(p)
		print(Aggre)
		Neuro = int(q)+int(r)+int(s)+int(t)
		print(Neuro)

		Openness = ((Open/8)*100)
		Conscientiousness = ((Cons/8)*100)
		Extraversion = ((Extra/8)*100)
		Aggreeableness = ((Aggre/8)*100)
		Neuroticism = ((Neuro/8)*100)
		'''
		traindata =pd.read_csv('train dataset.csv')
		array = traindata.values

		df=pd.DataFrame(array)
		maindf =df[[2,3,4,5,6]]
		mainarray=maindf.values
		#print(mainarray)

		temp=df[7]
		train_y =temp.values
		# print(train_y)
		# print(mainarray)

		for i in range(len(train_y)):
			train_y[i] =str(train_y[i])

		classifier = linear_model.LogisticRegression(multi_class='multinomial', solver='newton-cg',max_iter =1000)
		classifier.fit(mainarray, train_y)

		with open('classifier.pk1', 'wb') as pickle_file:
    		pickle.dump(classifier , pickle_file)

		'''
		with open('D:\\INTERNSHIP\\DJANGO\\personality_prediction\\app\\classifier.pk1', 'rb') as pickle_file:
			classifierpk = pickle.load(pickle_file)
        	       
		y_pred = classifierpk.predict([[Open,Cons,Extra,Aggre,Neuro]])

		return render(request,"score.html",{'tfmax':tfmax,'y_pred':y_pred,'Openness':Openness,'Conscientiousness':Conscientiousness,'Extraversion':Extraversion,'Aggreeableness':Aggreeableness,'Neuroticism':Neuroticism})
	
	return render(request,"test.html")

'''
def login(request):
	global rname,rpassword
	if request.method=='POST':
		rname=request.POST.get('rname')
		rpassword=request.POST.get('rpassword')
		if rname=='admin' and rpassword=='admin':
			return render(request, 'recruiter.html')
		else:
			return render(request, 'login.html')

	return render(request, 'login.html')


def recruiter(request):
	
	return render(request,"recruiter.html")
'''